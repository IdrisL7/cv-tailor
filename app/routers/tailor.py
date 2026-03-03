import uuid
import traceback
from pathlib import Path
from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from fastapi.responses import FileResponse, JSONResponse
from typing import Optional

from app.config import OUTPUT_DIR
from app.services.job_parser import fetch_job_description
from app.services.docx_handler import (
    read_docx,
    extract_sections,
    apply_tailored_content,
    save_docx,
    sections_to_text,
)
from app.services.ai_engine import (
    extract_keywords,
    run_pipeline,
)

# PDF support is optional (PyMuPDF may not be available on all platforms)
try:
    from app.services.pdf_handler import extract_sections_from_pdf
    PDF_SUPPORTED = True
except ImportError:
    PDF_SUPPORTED = False

router = APIRouter(prefix="/api", tags=["tailor"])


@router.post("/tailor")
async def tailor_cv(
    cv_file: UploadFile = File(...),
    job_url: Optional[str] = Form(None),
    job_text: Optional[str] = Form(None),
):
    try:
        if not job_url and not job_text:
            raise HTTPException(400, "Provide either a job URL or job description text.")

        filename = cv_file.filename or ""
        is_docx = filename.lower().endswith(".docx")
        is_pdf = filename.lower().endswith(".pdf")
        if not is_docx and not is_pdf:
            raise HTTPException(400, "Only .docx and .pdf files are supported.")
        if is_pdf and not PDF_SUPPORTED:
            raise HTTPException(400, "PDF support is not available on this server. Please upload a DOCX file instead.")

        # 1. Get job description
        jd_text = ""
        if job_url and job_url.strip():
            try:
                jd_text = await fetch_job_description(job_url.strip())
            except Exception as e:
                if not job_text:
                    raise HTTPException(
                        422,
                        f"Could not fetch job URL: {str(e)}. Please paste the job description text instead.",
                    )
        if not jd_text and job_text:
            jd_text = job_text.strip()
        if not jd_text:
            raise HTTPException(400, "No job description text available.")

        # 2. Save uploaded file and parse
        session_id = uuid.uuid4().hex[:12]
        ext = ".docx" if is_docx else ".pdf"
        input_path = OUTPUT_DIR / f"{session_id}_original{ext}"
        content = await cv_file.read()
        input_path.write_bytes(content)

        if is_docx:
            doc = read_docx(input_path)
            sections = extract_sections(doc)
        else:
            sections = extract_sections_from_pdf(input_path)
            doc = None

        # 3. AI pipeline (keyword extraction first, then tailor + prep in parallel)
        job_keywords = await extract_keywords(jd_text)
        tailored_sections, prep_summary = await run_pipeline(sections, job_keywords)

        # 4. Generate output file
        if is_docx and doc:
            doc = apply_tailored_content(doc, sections, tailored_sections)
            output_filename = f"{session_id}_tailored.docx"
            save_docx(doc, OUTPUT_DIR / output_filename)
        else:
            # PDF input: generate a new DOCX with tailored content
            output_filename = f"{session_id}_tailored.docx"
            _build_docx_from_sections(sections, tailored_sections, OUTPUT_DIR / output_filename)

        # 5. Keyword analysis
        cv_full_text = " ".join(" ".join(s.content_lines) for s in sections).lower()
        all_keywords = job_keywords.get("keywords", [])
        matched = [k for k in all_keywords if k.lower() in cv_full_text]
        missing = [k for k in all_keywords if k.lower() not in cv_full_text]

        return {
            "tailored_cv_filename": output_filename,
            "prep_summary": prep_summary,
            "keywords_matched": matched,
            "keywords_missing": missing,
            "job_title": job_keywords.get("title", ""),
            "company": job_keywords.get("company", ""),
        }
    except HTTPException:
        raise
    except Exception as e:
        traceback.print_exc()
        return JSONResponse(
            status_code=500,
            content={"detail": f"Internal error: {str(e)}"},
        )


@router.get("/download/{filename}")
async def download_file(filename: str):
    if ".." in filename or "/" in filename:
        raise HTTPException(400, "Invalid filename")
    file_path = OUTPUT_DIR / filename
    if not file_path.exists():
        raise HTTPException(404, "File not found")
    return FileResponse(
        path=str(file_path),
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def _build_docx_from_sections(
    sections: list,
    tailored_sections: dict[str, list[str]],
    output_path: Path,
) -> None:
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    for section in sections:
        if section.heading == "__HEADER__":
            for line in section.content_lines:
                p = doc.add_paragraph(line)
                for run in p.runs:
                    run.font.size = Pt(11)
            continue

        doc.add_heading(section.heading, level=2)

        lines = tailored_sections.get(section.heading, section.content_lines)
        for line in lines:
            # Only use bullet style for lines that look like bullet points
            is_bullet = line.lstrip().startswith(("•", "-", "–", "▪", "●", "○", "■"))
            style = "List Bullet" if is_bullet else "Normal"
            if is_bullet:
                line = line.lstrip("•-–▪●○■ ")
            p = doc.add_paragraph(line, style=style)
            for run in p.runs:
                run.font.size = Pt(11)

    doc.save(str(output_path))
