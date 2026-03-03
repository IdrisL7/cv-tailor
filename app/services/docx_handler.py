from docx import Document
from docx.text.paragraph import Paragraph
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional


@dataclass
class CVSection:
    heading: str
    heading_paragraph_index: int
    content_lines: list[str] = field(default_factory=list)
    paragraph_indices: list[int] = field(default_factory=list)


def extract_sections(doc: Document) -> list[CVSection]:
    sections: list[CVSection] = []
    current_section: Optional[CVSection] = None

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            if current_section:
                current_section.paragraph_indices.append(i)
            continue

        is_heading = (
            "Heading" in (para.style.name or "")
            or (len(text) < 60 and _is_visual_heading(para))
        )

        if is_heading:
            if current_section:
                sections.append(current_section)
            current_section = CVSection(
                heading=text,
                heading_paragraph_index=i,
                content_lines=[],
                paragraph_indices=[],
            )
        elif current_section:
            current_section.content_lines.append(text)
            current_section.paragraph_indices.append(i)
        else:
            current_section = CVSection(
                heading="__HEADER__",
                heading_paragraph_index=i,
                content_lines=[text],
                paragraph_indices=[i],
            )

    if current_section:
        sections.append(current_section)

    return sections


def _is_visual_heading(para: Paragraph) -> bool:
    if not para.runs:
        return False
    first_run = para.runs[0]
    is_bold = first_run.bold is True
    is_upper = para.text.strip().isupper()
    has_large_font = first_run.font.size is not None and first_run.font.size.pt >= 13
    return is_bold or is_upper or has_large_font


def replace_paragraph_text(paragraph: Paragraph, new_text: str) -> None:
    runs = paragraph.runs
    if not runs:
        return
    if len(runs) == 1:
        runs[0].text = new_text
    else:
        runs[0].text = new_text
        for run in runs[1:]:
            run.text = ""


def apply_tailored_content(
    doc: Document,
    sections: list[CVSection],
    tailored_sections: dict[str, list[str]],
) -> Document:
    for section in sections:
        if section.heading == "__HEADER__":
            continue
        key = section.heading
        if key not in tailored_sections:
            continue

        new_lines = tailored_sections[key]
        content_para_indices = [
            idx
            for idx in section.paragraph_indices
            if doc.paragraphs[idx].text.strip()
        ]

        for i, para_idx in enumerate(content_para_indices):
            para = doc.paragraphs[para_idx]
            if i < len(new_lines):
                replace_paragraph_text(para, new_lines[i])

    return doc


def read_docx(file_path: Path) -> Document:
    return Document(str(file_path))


def save_docx(doc: Document, output_path: Path) -> Path:
    doc.save(str(output_path))
    return output_path


def sections_to_text(sections: list[CVSection]) -> str:
    parts = []
    for s in sections:
        if s.heading != "__HEADER__":
            parts.append(f"## {s.heading}")
        parts.extend(s.content_lines)
        parts.append("")
    return "\n".join(parts)
